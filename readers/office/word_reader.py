"""
Word 文档读取器
"""
from pathlib import Path

from core.file_reader import BaseFileReader, FileContent
from config import CHUNK_SIZE, CHUNK_OVERLAP


class WordReader(BaseFileReader):
    """Word 文档读取器 (.docx)"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = [".docx", ".doc"]
    
    def read(self, file_path: Path) -> FileContent:
        """读取 Word 文档"""
        file_path = Path(file_path)
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            content = "\n\n".join(paragraphs)
            
            # 提取标题（使用第一个段落或文件名）
            title = doc.paragraphs[0].text if doc.paragraphs and doc.paragraphs[0].text else file_path.stem
            
        except Exception as e:
            # 如果 python-docx 失败，尝试作为纯文本读取
            content = f"[Error reading Word document: {e}]"
            title = file_path.stem
        
        # 提取元数据
        metadata = self.extract_metadata_from_path(file_path)
        
        # 分块
        chunks = self.chunk_content(content, CHUNK_SIZE, CHUNK_OVERLAP)
        
        return FileContent(
            source=str(file_path.absolute()),
            content=content,
            file_type="word",
            title=title[:100],  # 限制标题长度
            metadata=metadata,
            chunks=chunks,
            file_hash=self.calculate_hash(content)
        )
